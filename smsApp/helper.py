from .models import *
from django.db.models import Sum
from django.contrib.auth import get_user_model
from rest_framework.response import Response
from rest_framework import status
from django.db.models import Sum, Q
from orders.models import Order, OrderPayment
import requests
import json
from datetime import datetime
from django.utils.timezone import make_aware
from decouple import config
from companies.models import CompanySetting
import pymupdf
import fitz
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt
import os
import re


def sendSMSData(phone_numbers, message, is_scheduled, scheduled_time, user, send_type):
        message = message.encode().decode('UTF-8')
        sms_cost = 50
        general_setting = CompanySetting.objects.filter(
            company_setting=user.user_branch.company, setting_key='sms_unit_cost').first()
        if general_setting:
             sms_cost = general_setting.setting_value

        sending_mode = 'bulk' if len(phone_numbers) > 1 else 'single'

        for phone_number in phone_numbers:
                phone_number = f'256{phone_number[-9:]}'
                smsData = {
                    'user': user,
                    'send_type': send_type,
                    'sms_cost': sms_cost,
                    'phone_number': phone_number,
                    'message': message,
                    "is_scheduled": is_scheduled,
                    "scheduled_time": scheduled_time,
                    'sending_mode': sending_mode
                }

                if is_scheduled == True:
                    schedule_sms(smsData)
                if is_scheduled == False:
                    send_sms(smsData)

        # return Response({"message":""})


def sendBulkSMSData(recipients, message, is_scheduled, scheduled_time, user, send_type):
        message = message.encode().decode('UTF-8')
        sms_cost = 50
        general_setting = CompanySetting.objects.filter(
            company_setting=user.user_branch.company, setting_key='sms_unit_cost').first()
        if general_setting:
             sms_cost = general_setting.setting_value

        sending_mode = 'bulk' if len(recipients) > 1 else 'single'

        for recipient in recipients:
                phone_number = str(recipient['phone_number'])
                phone_number = f'256{phone_number[-9:]}'
                contact = User.objects.filter(
                    Q(phone_number__icontains=phone_number[-9:])).first()

                if contact is None:
                    contact = User.objects.create(
                        username=phone_number, first_name=recipient['name'], user_branch=user.user_branch, gender='O', phone_number=phone_number, user_type='Contact', user_added_by=user.id)

                smsData = {
                    'user': user,
                    'send_type': send_type,
                    'sms_cost': sms_cost,
                    'phone_number': phone_number,
                    'message': message,
                    "is_scheduled": is_scheduled,
                    "scheduled_time": scheduled_time,
                    'sending_mode': sending_mode
                }

                if is_scheduled == True:
                    schedule_sms(smsData)
                if is_scheduled == False:
                    send_sms(smsData)


def send_sms(smsData):
        message = smsData['message']
        sms_cost = smsData['sms_cost']
        phone_number = str(smsData['phone_number'])
        user = smsData['user']
        is_scheduled = smsData['is_scheduled']
        scheduled_time = smsData['scheduled_time']
        sending_mode = smsData['sending_mode']
        contact = User.objects.filter(
            Q(phone_number__icontains=phone_number[-9:])).first()
        sms_tran = None

        if contact is None:
            contact = User.objects.create(username=phone_number, first_name=phone_number, user_branch=user.user_branch,
                                          gender='O', phone_number=phone_number, user_type='Contact', user_added_by=user.id)

        if contact:
            if is_scheduled:
                sms_tran = UserSms.objects.filter(
                    recieved_by=contact, message=message, is_scheduled=is_scheduled).first()
            else:
                # send sms
                sms_tran = UserSms.objects.create(is_scheduled=is_scheduled, scheduled_time=scheduled_time, sms_cost=sms_cost, message=message, provider=config(
                    'ACTIVE_SMS_PROVIDER'), sending_mode=sending_mode, telephone=phone_number, status='pending', company=user.user_branch.company, sent_by=user, recieved_by=contact)

        if sms_tran:
            if config('ACTIVE_SMS_PROVIDER') == 'EGO':
                data = {'method': 'SendSms',
                        'userdata': {
                            'username': config('EGO_USER_NAME'),
                            'password': config('EGO_PASSWORD')
                        },
                        'msgdata': [{'number': phone_number, 'message': message, 'senderid': 'Egosms'}]}

                send_sms = requests.post(config('EGO_SMS_URL'), json = data)

                if send_sms.status_code == 200 or send_sms.status_code == '200':
                    send_resp = json.loads(send_sms.text)
                    if send_resp['Status'] == 'OK':
                        sms_tran.status = 'sent'
                        sms_tran.save()
                    else:
                        sms_tran.status = 'failed'
                        sms_tran.save()
                else:
                    sms_tran.status = 'failed'
                    sms_tran.save() 


def schedule_sms(smsData):
        message = smsData['message']
        sms_cost = smsData['sms_cost']
        phone_number = str(smsData['phone_number'])
        user = smsData['user']
        is_scheduled = smsData['is_scheduled']
        scheduled_time = smsData['scheduled_time']
        sending_mode = smsData['sending_mode']
        contact = User.objects.filter(
            Q(phone_number__icontains=phone_number[-9:])).first()

        if contact is None:
            contact = User.objects.create(username=phone_number, first_name=phone_number, user_branch=user.user_branch,
                                          gender='O', phone_number=phone_number, user_type='Contact', user_added_by=user.id)

        if contact:
            UserSms.objects.create(is_scheduled=is_scheduled, scheduled_time=scheduled_time, sms_cost=sms_cost, message=message, provider=config(
                'ACTIVE_SMS_PROVIDER'), sending_mode=sending_mode, telephone=phone_number, status='pending', company=user.user_branch.company, sent_by=user, recieved_by=contact)


def send_schuduled_sms():
    time = timezone.now()
    sms_trans = UserSms.objects.filter(is_scheduled=True,status='pending',scheduled_time__lte=time)
    for sms_tran in sms_trans:
        smsData = {
                    'user':sms_tran.sent_by,
                    'send_type':sms_tran.sending_mode,
                    'sms_cost':sms_tran.sms_cost,
                    'phone_number':sms_tran.recieved_by.phone_number,
                    'message':sms_tran.message,
                    "is_scheduled":sms_tran.is_scheduled,
                    "scheduled_time":sms_tran.scheduled_time,
                    'sending_mode':sms_tran.sending_mode
                }
        send_sms(smsData)

            
def bulk_group_contacts_update(group_id,contacts,user):
    for contact in contacts:
        contact_id = contact['id']
        if contact_id:
            assigned_contact = ContactAssignedGroup.objects.filter(contact_group__id=group_id,contact__id=contact_id).first()
            if assigned_contact:  
                assigned_contact.is_active=True
                assigned_contact.save()
            else:
                ContactAssignedGroup.objects.create(contact_group_id=group_id,contact_id=contact_id,added_by=user) 
    

def trash_update_orders():
    requests = SmsRequest.objects.all()
    if requests:  
        for request in requests:
            Order.objects.filter(ext_ref_no=request.id,trans_type='sms-top-up').delete()
            saved_order = Order.objects.create(order_no=generate_invoice_no(request.company.id),ext_ref_no=request.id,company=request.company,trans_type='sms-top-up', customer=request.requested_by, date_added=request.date_added)
            if saved_order:
                OrderPayment.objects.create(payment_method=request.payment_method,naration="Sms top up",amount=request.approved_amount,ref_no=generate_ref_no(request.company.id),order=saved_order,added_by=request.requested_by, date_added=request.date_added)



def get_account_sms_balance(company):
    purchases = SmsRequest.objects.filter(status='approved', company=company).aggregate(total_amount = Sum('approved_amount') )
    sms_transactions = UserSms.objects.filter(company=company, status='sent').aggregate(total_sms_cost = Sum('sms_cost') )
    sms_awards = CompanyFreeSmsAward.objects.filter(company=company, status='approved').aggregate(total_awards = Sum('amount') )
   
    total_awards = 0
    total_payments  = 0
    total_sms_transactions = 0
    
    if sms_awards:
        total_awards = sms_awards['total_awards'] if sms_awards['total_awards'] != None else 0
    
    if purchases:
        total_payments = purchases['total_amount'] if purchases['total_amount'] != None else 0
    
    if sms_transactions:
        total_sms_transactions = sms_transactions['total_sms_cost'] if sms_transactions['total_sms_cost'] != None else 0

    balance = (float(total_awards) + float(total_payments)) - float(total_sms_transactions) 
    return balance if balance > 0 else 0

def pad_reference_number(number,max=5):
    number = str(number)
    leading_zeros = max - len(number)
    if leading_zeros > 0:
        number = '0' * leading_zeros + number
    
    return number

def generate_invoice_no(company_id, prefix = None):
    base = 'Horbtec_INV'

    last_reference = Order.objects.filter(company__id=company_id,order_no__icontains=base).order_by('-id')
    new_ref = 1

    if(len(last_reference) > 0):
        new_ref = int(last_reference[0].order_no.replace(base,'')) + 1

    return base + pad_reference_number(new_ref,6)


def generate_ref_no(company_id, prefix = None):
    now = datetime.now()
    base = f'P{now.strftime('%d%b%y')}RE'
    new_ref = 1
    last_reference = OrderPayment.objects.filter(order__company__id=company_id,ref_no__icontains=base).order_by('-id')
    if(len(last_reference) > 0):
        new_ref = int(last_reference[0].ref_no.replace(base,'')) + 1

    return base + pad_reference_number(new_ref,3)


def save_request_order(sms_request):
        if sms_request:  
            update_existing_request_orders(sms_request.company.id)
            ''' saved_order = Order.objects.create(order_no=generate_invoice_no(sms_request.company.id),ext_ref_no=sms_request.id,company=sms_request.company,trans_type='sms-top-up', customer=sms_request.requested_by,added_by=sms_request.requested_by)
            if saved_order:
                OrderPayment.objects.create(payment_method=sms_request.payment_method,naration="Sms top up",amount=sms_request.approved_amount,ref_no=generate_ref_no(sms_request.company.id),order=saved_order,trans_type='sms-top-up',added_by=sms_request.requested_by) '''

def update_existing_request_orders(company_id):
        requests = SmsRequest.objects.filter(company__id = company_id).order_by('id')
        if requests:  
            for request in requests:
                order = Order.objects.filter(ext_ref_no=request.id,trans_type='sms-top-up').order_by('id')
                if not order:
                    saved_order = Order.objects.create(order_no=generate_invoice_no(request.company.id),ext_ref_no=request.id,company=request.company,trans_type='sms-top-up', customer=request.requested_by)
                    if saved_order:
                        OrderPayment.objects.create(payment_method=request.payment_method,naration="Sms top up",amount=request.approved_amount,ref_no=generate_ref_no(request.company.id),order=saved_order,added_by=request.requested_by)

def search_pdf_and_get_rect_dimensions(pdf_path, search_text):
    # Open the PDF document
    doc = fitz.open(pdf_path)
    rect_dimensions = {'x1':None,'x1':None,'x2':None,'y2':None,'page_index':None}
    try:
      
        # Iterate through each page in the document
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Search for the text on the current page
            text_instances = page.search_for(search_text)
        
            # Iterate through each text instance found
            for rect in text_instances:
                # Get the width of the rectangle dimensions containing the text
                rect_dimensions = {'x1':rect[0],'x2':rect[2],'y1':rect[1],'y2':rect[3],'page_index':page_num}
                # Print the page number, rectangle coordinates, and width
                print(f"Page {page_num + 1}: Text found at {rect}, dimensions: {rect_dimensions}")
    
    finally:
        doc.close()
    return rect_dimensions


def insert_html_at_position(directory_path,pdf_path,pdf_out_put_path,search_text, html_content,user):
   
    file_path = directory_path + 'munites.pdf'
    file_path2 = directory_path + 'munites_final.pdf'
    doc = fitz.open(file_path)
    rect_dimensions = search_pdf_and_get_rect_dimensions(file_path, search_text)
    if rect_dimensions['x1']:
        page_index = rect_dimensions['page_index']
        page = doc[page_index]
        rect = (round(rect_dimensions['x1']),round(rect_dimensions['y1']) + 5, rect_dimensions['x2'] - rect_dimensions['x1'], 500)
        # Create an HTML annotation (box) at the specified position
        page.insert_htmlbox(rect, '<p>Hello how are you, <br/> cool stufff')  # place into the rectangle
    doc.save(file_path2)
    doc.close()
    FileObject.objects.create(**{
        "title": pdf_out_put_path,
        "description": file_path2,
        "url_path": file_path2,
        "user":user,
        "attachment_type": 'image-to-text'
    })

def image_to_string(directory_path,image_paths,user):
    
    if len(image_paths) > 0:
        doc = Document()
        text1 = pytesseract.image_to_string(directory_path + image_paths[0], lang='eng', config='--psm 6')
        text1 = re.sub(r'\s+', '', text1)
        file_name = f'{text1[:20]}.docx'
        for image_path in image_paths:
            # Perform OCR with specified language and custom configurations
            text = pytesseract.image_to_string(directory_path + image_path, lang='eng', config='--psm 6')
            # Create a new Word document
            # Add formatted text
            paragraph = doc.add_paragraph()

            # Split text into lines
            lines = text.split('\n')
            
            for line in lines:
                run = paragraph.add_run(line)
                run.font.size = Pt(12)
                run.font.name = 'Arial'
                # Add newline after each line
                paragraph.add_run('\n')  
            try:
                os.remove(directory_path + image_path)
            except FileNotFoundError:
                print(f"File {directory_path + image_path} not found.")

        # Save the document
        doc.save(directory_path + file_name)
        FileObject.objects.create(**{
            "title": file_name,
            "description": file_name,
            "url_path": directory_path + file_name,
            "user":user,
            "attachment_type": 'image-to-text'
        })